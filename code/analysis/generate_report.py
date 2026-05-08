#!/usr/bin/env python3
"""
Research Report Generator: LLM NLP Signal Impact on DRL Trading Performance
Mak Chi Lam | 3035279753 | STAT8307

Uses python-docx for production-quality DOCX output.
Run with: /opt/anaconda3/bin/python3 generate_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ── Color Palette ──
NAVY = RGBColor(0, 51, 102)
TEAL = RGBColor(8, 145, 178)
CORAL = RGBColor(220, 80, 60)
DARK = RGBColor(33, 37, 41)
GRAY = RGBColor(108, 117, 125)
WHITE = RGBColor(255, 255, 255)
LIGHT_BG = RGBColor(248, 249, 250)

# ── Document Setup ──
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Normal style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for lvl in range(1, 5):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = NAVY
    hs.paragraph_format.space_before = Pt(18 if lvl <= 2 else 12)
    hs.paragraph_format.space_after = Pt(6)
    if lvl == 1:
        hs.font.size = Pt(18)
    elif lvl == 2:
        hs.font.size = Pt(14)
    elif lvl == 3:
        hs.font.size = Pt(12)
    else:
        hs.font.size = Pt(11)

# Table counter
tbl_num = [0]

# ── Helper Functions ──

def add_para(text, bold=False, italic=False, size=11, color=None, align=None, space_after=6, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_rich_para(parts, align=None, space_after=6, indent=None):
    """parts = list of (text, bold, italic, size, color)"""
    p = doc.add_paragraph()
    for text, bold, italic, size, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_table(headers, rows, col_widths=None, caption=None, header_color=NAVY):
    tbl_num[0] += 1
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Cell shading
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): str(header_color)
        })
        shading.append(shd)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = t.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternate row shading
            if row_idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F0F4F8'
                })
                shading.append(shd)

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # Caption
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f'Table {tbl_num[0]}: ')
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = NAVY
        run2 = p.add_run(caption)
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY
        p.paragraph_format.space_before = Pt(4)

    return t

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
add_para('Does Better NLP Signal Quality\nImprove DRL Trading Performance?', bold=True, size=22, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para('Comparing LLM Intelligence in Financial Sentiment Analysis\nfor Trading Decision Support', size=14, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('─' * 40, size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('Mak Chi Lam', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Student ID: 3035279753', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('STAT8307 Natural Language Processing and Text Analytics', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
add_para('May 2026', size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Abstract', level=1)

add_para(
    'This study investigates whether the quality of NLP signals generated by different large language models (LLMs) '
    'affects the trading performance of a deep reinforcement learning (DRL) agent. We compare four LLM configurations '
    '-- Gemini 3.1 Pro, Qwen3.5-27B (base), Qwen3.5-27B with QLoRA fine-tuning, and Qwen3.5-27B with QA-LoRA '
    'fine-tuning -- as signal generators within the PrimoGPT/PrimoRL framework adapted from Botunac (2025). '
    'Despite substantial differences in signal quality (all-zero-day rates ranging from 6.9% to 39.9%, per-cell '
    'zero-fill from 32% to 91%, and 66-96% of non-zero signal cells differing across LLMs), a one-way ANOVA across '
    '2,800+ training runs finds no significant effect of LLM choice on trading returns (F=0.034, p=0.998, '
    '\u03b7\u00b2=0.006). Even the pooled comparison of Has-NLP versus No-NLP agents yields a negligible effect '
    '(+0.81 percentage points, p=0.736, Cohen\'s d=0.037). However, an ablation study reveals a paradox: '
    'NLP-only agents significantly outperform Tech-only agents (+1.8pp, p<0.0001, d=0.36). We resolve this '
    'by showing that NLP and technical indicators encode similar information (redundancy), but NLP delivers it in a '
    'simpler form -- 7 less-correlated features yielding 97.3% active trading versus 8 correlated technical features '
    'yielding only 69.6%. The null result reflects information redundancy, not signal irrelevance. For practitioners, '
    'this implies that investing in better LLMs for trading signal generation yields zero marginal return; instead, '
    'state space design -- fewer, less correlated features -- may be more impactful than signal quality.'
)

add_rich_para([
    ('Keywords: ', True, False, 11, NAVY),
    ('deep reinforcement learning, large language models, financial sentiment analysis, trading signals, PPO, ablation study, information redundancy', False, True, 11, GRAY),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1.  Introduction',
    '    1.1  Research Question',
    '    1.2  Reference Study and Research Gap',
    '    1.3  Contributions',
    '2.  System Architecture',
    '    2.1  PrimoGPT: NLP Signal Generation',
    '    2.2  PrimoRL: DRL Trading Environment',
    '    2.3  LLM Configurations',
    '3.  Methodology',
    '    3.1  Two-Stage Experimental Design',
    '    3.2  Controls and Bias Prevention',
    '    3.3  Statistical Framework',
    '4.  Challenges and Iterative Refinement',
    '    4.1  Look-Ahead Bias',
    '    4.2  Unrealistic Reported Results',
    '    4.3  Failed Replication',
    '    4.4  Model Variance and Do-Nothing Problem',
    '5.  Experiment Evolution (v1-v8)',
    '    5.1  v1-v2: Raw Observations',
    '    5.2  v3: Selective Normalization',
    '    5.3  v4: Ensemble Aggregation',
    '    5.4  v5: Differential Sharpe Reward',
    '    5.5  v6: Hyperparameter Tuning',
    '    5.6  v7: 30-Seed Definitive Experiment',
    '    5.7  v8: Ablation Study',
    '6.  Results',
    '    6.1  Stage 1: NLP Signal Quality',
    '    6.2  Stage 2: Trading Performance (v7)',
    '    6.3  Feature Importance Analysis',
    '    6.4  Ablation Results (v8)',
    '7.  Discussion',
    '    7.1  Why NLP Signals Do Not Improve Trading',
    '    7.2  The Ablation Paradox and Resolution',
    '    7.3  Ensemble Granularity',
    '    7.4  Alternative Algorithms',
    '8.  Conclusion',
    '9.  References',
    'A.  Appendix: Hyperparameters and Seed List',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1.2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Research Question', level=2)

add_para(
    'The integration of large language models (LLMs) into financial trading systems has generated considerable '
    'excitement, driven by the premise that better language understanding should yield better trading decisions. '
    'Financial sentiment analysis -- extracting structured trading signals from unstructured news text -- has emerged '
    'as a key application domain. However, a fundamental question remains unanswered: does the quality of LLM-generated '
    'NLP signals actually translate into improved trading performance?'
)

add_para(
    'This study addresses that question directly. We test whether signals generated by different LLM configurations '
    '-- varying in model architecture, parameter count, and fine-tuning strategy -- produce measurably different '
    'outcomes when fed to a deep reinforcement learning (DRL) trading agent. Specifically, we compare a top-tier '
    'proprietary model (Gemini 3.1 Pro) against three configurations of an open-weight model (Qwen3.5-27B: base, '
    'QLoRA-tuned, and QA-LoRA-tuned) within the PrimoGPT/PrimoRL framework.'
)

doc.add_heading('1.2 Reference Study and Research Gap', level=2)

add_para(
    'This project is based on the article "Automated Trading Framework Using LLM-Driven Features and Deep '
    'Reinforcement Learning" by Botunac (2025), published in the journal Big Data and Cognitive Computing. '
    'That study introduced the PrimoGPT system, which uses an LLM to generate 7 financial sentiment signals '
    'from news articles, feeds them into a PPO-based trading agent (PrimoRL), and reported impressive single-stock '
    'results: +58.47% return on Netflix with a Sharpe ratio of 2.81.'
)

add_para(
    'However, the reference study applied only the QA-LoRA-tuned model and did not compare different LLMs. '
    'This leaves two critical gaps: (1) whether different LLMs generate signals of different quality, and '
    '(2) whether such quality differences propagate to downstream trading performance. Our study fills both gaps.'
)

add_para(
    'Additionally, during our replication attempts, we identified several methodological concerns in the reference '
    'study: a look-ahead bias in NLP signal alignment, the absence of transaction costs, and results reported '
    'from only a single random seed. These concerns motivated a systematic redesign of the experimental methodology.'
)

doc.add_heading('1.3 Contributions', level=2)

add_bullet('A systematic comparison of 4 LLM configurations as NLP signal generators, benchmarking signal quality '
           '(directional accuracy, information coefficient, zero-fill rates) and downstream trading performance.', bold_prefix='Signal quality benchmarking: ')
add_bullet('A 30-seed definitive experiment (750 runs) with one-way ANOVA demonstrating that LLM choice has no '
           'significant effect on DRL trading performance (p=0.998, \u03b7\u00b2=0.006).', bold_prefix='Null result with rigorous evidence: ')
add_bullet('An ablation study revealing that NLP-only agents outperform Tech-only agents (+1.8pp, p<0.0001), '
           'which we resolve through a simpler-state-space explanation: NLP and tech encode similar information, '
           'but NLP delivers it in a simpler form (97.3% vs 69.6% active rate).', bold_prefix='Ablation paradox and resolution: ')
add_bullet('Iterative methodological refinement across 8 experimental versions and 2,800+ training runs, '
           'addressing look-ahead bias, feature scale mismatch, reward function design, and model variance.', bold_prefix='Methodological improvements: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. System Architecture', level=1)

add_para(
    'The system, adapted from Botunac (2025), consists of two modules: PrimoGPT (NLP signal generation) and '
    'PrimoRL (DRL trading). Figure 1 illustrates the data flow.'
)

# Flow diagram as table
flow = doc.add_table(rows=2, cols=1, style='Table Grid')
flow.alignment = WD_TABLE_ALIGNMENT.CENTER

cell0 = flow.rows[0].cells[0]
cell0.text = ''
p0 = cell0.paragraphs[0]
run0 = p0.add_run('PrimoGPT (NLP Module)')
run0.bold = True
run0.font.color.rgb = TEAL
run0.font.size = Pt(10)
p0b = cell0.add_paragraph()
run0b = p0b.add_run(
    'Financial News \u2192 LLM \u2192 7 NLP Signals\n'
    '[Sentiment, Price Impact, Trend Direction, Earnings Impact,\n'
    ' Investor Confidence, Risk Profile Change, News Relevance]'
)
run0b.font.name = 'Consolas'
run0b.font.size = Pt(9)

cell1 = flow.rows[1].cells[0]
cell1.text = ''
p1 = cell1.paragraphs[0]
run1 = p1.add_run('PrimoRL (DRL Module)')
run1.bold = True
run1.font.color.rgb = NAVY
run1.font.size = Pt(10)
p1b = cell1.add_paragraph()
run1b = p1b.add_run(
    '18-dim State: [Cash, Price, Holdings, 8 Tech, 7 NLP] \u2192 PPO \u2192 Action [-1, +1]\n'
    'Tech: MACD, Boll_UB, Boll_LB, RSI_30, CCI_30, DX_30, Close_30_SMA, Close_60_SMA'
)
run1b.font.name = 'Consolas'
run1b.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Figure 1: ')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = NAVY
run2 = p.add_run('System architecture: PrimoGPT generates NLP signals from financial news; PrimoRL combines them with market data for PPO-based trading.')
run2.italic = True
run2.font.size = Pt(9)
run2.font.color.rgb = GRAY

doc.add_heading('2.1 PrimoGPT: NLP Signal Generation', level=2)

add_para(
    'PrimoGPT processes daily financial news articles for each stock through an LLM, which outputs 7 continuous '
    'signal features in the range [0, 1]: Sentiment, Price Impact Potential, Trend Direction, Earnings Impact, '
    'Investor Confidence, Risk Profile Change, and News Relevance. These signals are aligned to trading days and '
    'shifted by one day (shift(1)) to prevent look-ahead bias: the agent observes day T-1\'s NLP signal when '
    'making day T\'s trading decision.'
)

doc.add_heading('2.2 PrimoRL: DRL Trading Environment', level=2)

add_para(
    'PrimoRL constructs an 18-dimensional state vector per stock comprising three feature groups: '
    '(1) portfolio state (cash balance, closing price, share holdings), '
    '(2) 8 technical indicators (MACD, Bollinger Bands upper/lower, RSI-30, CCI-30, DX-30, 30-day SMA, 60-day SMA), '
    'and (3) the 7 NLP signals from PrimoGPT. A PPO agent with continuous action space [-1, +1] determines '
    'daily buy/sell amounts. The environment applies 0.1% transaction costs per trade.'
)

add_table(
    ['Feature Group', 'Features', 'Dimension', 'Range'],
    [
        ['Portfolio State', 'Cash, Close Price, Holdings', '3', 'Cash ~100K, Price ~200'],
        ['Technical Indicators', 'MACD, Boll_UB/LB, RSI, CCI, DX, SMA_30/60', '8', 'Standardized'],
        ['NLP Signals', 'Sentiment, Impact, Trend, Earnings, Confidence, Risk, Relevance', '7', '[0, 1] raw; 25x scaled'],
    ],
    col_widths=[3.0, 5.5, 2.0, 4.0],
    caption='State vector composition (18 dimensions per stock)'
)

doc.add_heading('2.3 LLM Configurations', level=2)

add_para(
    'We compare four LLM configurations plus a No-NLP baseline. The configurations span a proprietary API model '
    'and three variants of an open-weight model with increasing levels of fine-tuning:'
)

add_table(
    ['Configuration', 'Base Model', 'Fine-Tuning', 'All-Zero-Day Rate', 'Per-Cell Zero-Fill'],
    [
        ['Gemini 3.1 Pro', 'Gemini 3.1 Pro', 'None (API)', '6.9%', '32.2%'],
        ['Qwen Base', 'Qwen3.5-27B', 'None', '39.9%', '57.9%'],
        ['QLoRA', 'Qwen3.5-27B', 'QLoRA (r=16)', '18.2%', '82.6%'],
        ['QA-LoRA', 'Qwen3.5-27B', 'QA-LoRA (r=16)', '12.0%', '90.9%'],
        ['No-NLP Baseline', '\u2014', '\u2014', '100%', '100%'],
    ],
    col_widths=[2.5, 2.5, 2.5, 3.0, 3.0],
    caption='LLM configurations and zero-fill characteristics'
)

add_para(
    'Two distinct zero-fill metrics are reported: "All-Zero-Day Rate" measures the percentage of trading days '
    'where all 7 signals are simultaneously zero, while "Per-Cell Zero-Fill" measures the percentage of individual '
    'signal-day cells that are zero. These differ dramatically: QA-LoRA has only 12% all-zero-day rate but 91% '
    'per-cell zero-fill, meaning most days have at least one non-zero signal but the vast majority of individual '
    'signals are zero. Across LLMs, 66-96% of non-zero signal cells differ, indicating fundamentally different '
    'signal distributions.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Methodology', level=1)

doc.add_heading('3.1 Two-Stage Experimental Design', level=2)

add_para(
    'Our methodology follows a two-stage design that separates signal quality measurement from trading validation:'
)

add_rich_para([
    ('Stage 1: NLP Signal Quality Benchmarking. ', True, False, 11, TEAL),
    ('We directly measure the predictive quality of NLP signals generated by each LLM, using directional accuracy '
     '(percentage of non-zero signals correctly predicting next-day return direction, with random baseline at 33% '
     'for three-class classification), information coefficient (Spearman rank correlation between signal value '
     'and next-day return), and zero-fill rates.', False, False, 11, DARK),
], indent=0.5, space_after=8)

add_rich_para([
    ('Stage 2: Trading Validation. ', True, False, 11, TEAL),
    ('A fixed PPO environment where the only variable is the NLP source. Same hyperparameters, same 30 seeds, '
     'same price data -- the only difference is which LLM generated the features. We also include a No-NLP baseline '
     'where all NLP features are set to zero.', False, False, 11, DARK),
], indent=0.5, space_after=8)

add_table(
    ['Parameter', 'Value'],
    [
        ['Algorithm', 'PPO (Proximal Policy Optimization)'],
        ['Architecture', 'Deep [128, 128, 64]'],
        ['Seeds', '30 per condition'],
        ['Tickers', 'AAPL, AMZN, CRM, MSFT, NFLX'],
        ['Training Period', '2022-04 to 2024-07'],
        ['Test Period', '2024-08 to 2025-02 (144 trading days)'],
        ['Transaction Costs', '0.1% per trade'],
        ['Learning Rate', '1e-4'],
        ['Batch Size', '256'],
        ['N-Steps', '2048'],
        ['Entropy Coefficient', '0.01'],
        ['Total Timesteps', '200,000'],
        ['Reward Function', 'Dollar delta (change in portfolio value)'],
        ['Normalization', 'Z-score on cash/price/holdings; 25x upscale on NLP'],
        ['Look-Ahead Prevention', 'shift(1) on all NLP columns'],
    ],
    col_widths=[4.5, 10.0],
    caption='Experimental hyperparameters for the definitive experiment (v7)'
)

doc.add_heading('3.2 Controls and Bias Prevention', level=2)

add_bullet('All NLP signal columns are shifted by one trading day, so the agent uses day T-1\'s NLP data when making day T\'s trading decision.', bold_prefix='Look-ahead bias prevention: ')
add_bullet('A realistic 0.1% transaction cost is applied to every trade, unlike the reference study which used zero transaction costs.', bold_prefix='Transaction costs: ')
add_bullet('30 pre-defined seeds are used per condition; all results are reported without cherry-picking. Agents that fail to converge (SVD errors) or produce degenerate strategies (Sharpe < 0 with > 50 trades) are documented.', bold_prefix='No cherry-picking: ')
add_bullet('The only difference between experimental conditions is the NLP signal source; all other variables (price data, technical indicators, hyperparameters, seeds) are held constant.', bold_prefix='Controlled comparison: ')
add_bullet('We acknowledge that our 5 test tickers are all large-cap US technology stocks that appreciated during the test period. This introduces survivorship and selection bias, but the within-study LLM comparison remains valid since all configurations face the same tickers.', bold_prefix='Selection bias acknowledgment: ')

doc.add_heading('3.3 Statistical Framework', level=2)

add_para(
    'We employ three levels of statistical analysis:'
)

add_rich_para([
    ('(1) One-way ANOVA ', True, False, 11, DARK),
    ('across the 5 LLM configurations (4 LLM + No-NLP) to test the global null hypothesis that LLM choice has '
     'no effect on mean returns. The test uses ensemble-level returns (n=5 per group) and reports F-statistic, '
     'p-value, and eta-squared (\u03b7\u00b2) as the effect size measure.', False, False, 11, DARK),
], indent=0.5, space_after=6)

add_rich_para([
    ('(2) Paired same-seed t-tests ', True, False, 11, DARK),
    ('at the individual agent level, matching agents by (ticker, seed) pair. Each comparison has n=95-100 paired '
     'observations. Raw p-values are reported; Benjamini-Hochberg (BH) correction is available as a robustness check '
     'but does not change the substantive conclusion.', False, False, 11, DARK),
], indent=0.5, space_after=6)

add_rich_para([
    ('(3) Pooled Has-NLP vs No-NLP comparison ', True, False, 11, DARK),
    ('aggregating across all LLM configurations to maximize statistical power for detecting any NLP effect. '
     'Effect size is measured by Cohen\'s d.', False, False, 11, DARK),
], indent=0.5, space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. CHALLENGES AND ITERATIVE REFINEMENT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Challenges and Iterative Refinement', level=1)

add_para(
    'Our investigation began as a replication of Botunac (2025) but encountered four major methodological '
    'challenges that necessitated systematic fixes across 8 experimental iterations.'
)

doc.add_heading('4.1 Look-Ahead Bias', level=2)

add_para(
    'The reference code did not apply shift(1) to NLP signal columns, meaning the agent could observe day T\'s '
    'news signals before making day T\'s trading decision. This constitutes a look-ahead bias that artificially '
    'inflates performance. We corrected this by applying shift(1) to all NLP features, ensuring that the agent '
    'only uses day T-1\'s NLP data when deciding day T\'s trade.'
)

add_para(
    'To quantify the bias effect, we ran controlled experiments comparing shifted vs. unshifted NLP signals. '
    'The effect varied by ticker (AAPL: -0.5pp, NFLX: +8.6pp) but was not statistically significant (both p>0.05), '
    'likely because NLP signals are slowly varying and the bimodal variance in trading outcomes swamps the effect.'
)

doc.add_heading('4.2 Unrealistic Reported Results', level=2)

add_para(
    'The reference study reported a +58.47% return on Netflix with a Sharpe ratio of 2.81 from a single-stock '
    'PPO agent with zero transaction costs. Such results are unrealistic for a DRL trading strategy: professional '
    'quantitative strategies targeting single stocks rarely achieve sustained Sharpe ratios above 2.0, and the '
    'absence of transaction costs further inflates returns. These impressive but unrealistic results were the '
    'initial motivation for replication.'
)

doc.add_heading('4.3 Failed Replication', level=2)

add_para(
    'We could not reproduce the reference results even using the original open-source code with their stated '
    'configuration. The paper reports results from a single random seed, and we suspect cherry-picking: individual '
    'seed results vary enormously (e.g., AAPL returns range from -2% to +15% across seeds), making single-seed '
    'results highly unrepresentative.'
)

doc.add_heading('4.4 Model Variance and Do-Nothing Problem', level=2)

add_para(
    'PPO is notoriously seed-dependent. In our 30-seed experiments, approximately 27% of trained agents converge '
    'to do-nothing policies (no trades or trivial positions), 2% produce oscillating strategies (many trades with '
    'negative Sharpe), and only 71% produce active, sensible trading behavior. This bimodal distribution -- '
    'where most agents either trade meaningfully or do nothing at all -- overwhelms any subtle signal quality '
    'effects. The implication is clear: single-seed results are meaningless for comparing LLM configurations; '
    'multi-seed studies are essential.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. EXPERIMENT EVOLUTION (v1-v8)
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Experiment Evolution (v1\u2013v8)', level=1)

add_para(
    'Across 8 experimental versions, each time we improved the methodology to give LLM a fairer chance of '
    'demonstrating its value, we obtained the same answer. Table 2 summarizes the evolution.'
)

add_table(
    ['Version', 'Key Change', 'Seeds', 'ANOVA p', 'Spread', 'Active Rate'],
    [
        ['v1', 'Raw observations (baseline)', '3', '0.999', '\u2014', '~54%'],
        ['v2', 'Additional reward variants', '3', '0.999', '\u2014', '~54%'],
        ['v3', 'Z-score + 25x NLP upscale', '3', '0.9999', '2.3pp', '~54%'],
        ['v4', 'Ensemble of active agents', '10', '1.000', '0.48pp', '90%'],
        ['v5', 'Differential Sharpe reward', '10', '0.9998', '1.18pp', '56%'],
        ['v6', '3 HP configs (Deep best)', '10', '>0.99', '1.4\u20135.5pp', '90% (Deep)'],
        ['v7', '30 seeds, Deep arch (definitive)', '30', '0.998', '4.9pp', '70.9%'],
        ['v8', 'Ablation (Full/Tech/NLP)', '30', '(paired t)', '1\u20132pp', '69\u201397%'],
    ],
    col_widths=[1.2, 4.5, 1.5, 2.0, 2.0, 2.5],
    caption='Experiment evolution across 8 versions. ANOVA p-values remain >0.99 throughout v1\u2013v7.'
)

doc.add_heading('5.1 v1\u2013v2: Raw Observations', level=2)

add_para(
    'The initial experiments used raw, unnormalized observations. The state vector contained cash around 100,000, '
    'price around 200, and NLP features in [-2, +2]. This extreme scale mismatch made NLP features literally '
    'invisible to gradient-based learning: permutation importance attributed only 0.8% of decision variance to NLP '
    '(compared to 83.9% for portfolio state and 15.3% for technical indicators). ANOVA yielded p=0.999. '
    'While the null was expected given the scale mismatch, v1 established the baseline and the feature importance '
    'decomposition that would prove crucial later.'
)

doc.add_heading('5.2 v3: Selective Normalization', level=2)

add_para(
    'Having identified the feature scale mismatch as the likely root cause of the v1/v2 null, v3 applied z-score '
    'normalization to cash, price, and holdings, plus a 25x upscale on NLP features, bringing them to '
    'approximately [-50, +50]. NLP was now visible to gradients. However, the ANOVA became even more null '
    '(p=0.9999). This was a critical finding: making NLP visible to the learning algorithm does NOT change '
    'the conclusion. The null result is genuine, not an artifact of scale.'
)

doc.add_heading('5.3 v4: Ensemble Aggregation', level=2)

add_para(
    'Version 4 increased from 3 to 10 seeds and introduced ensemble aggregation: averaging actions across '
    'active agents (filtering out do-nothing agents with |return| < 0.1%). The stunning result was that '
    'No-NLP actually won the ensemble comparison (+24.44% vs. Gemini +23.96%), with a spread of only 0.48pp '
    'between best and worst configurations. ANOVA gave F=0.0003, p=1.000, eta-squared=0.000056 -- '
    'a literally perfect null. The active rate improved to 90% thanks to better seed coverage.'
)

doc.add_heading('5.4 v5: Differential Sharpe Reward', level=2)

add_para(
    'To test whether the null was reward-function-specific, v5 replaced the dollar_delta reward with the '
    'differential Sharpe ratio. This required careful normalization: z-score normalization compresses dollar '
    'rewards to near-zero, causing do-nothing policies. The solution was to use NLP-only upscale (no z-score) '
    'with the diff_sharpe reward. Results were lower overall (~12% vs. ~24% with dollar_delta) and the do-nothing '
    'rate was higher (44% vs. 10%), but the ANOVA remained null at p=0.9998. The null result holds across '
    'reward functions.'
)

doc.add_heading('5.5 v6: Hyperparameter Tuning', level=2)

add_para(
    'Version 6 tested three architecture configurations: Default [64, 64], Deep [128, 128, 64], and Explore '
    '[64, 64] with high entropy. The Deep architecture achieved the highest active rate (90% vs. 54% Default, '
    '64% Explore) and was selected for the definitive experiment. However, no architecture made LLM significant '
    '(all ANOVA p>0.99), and the choice of architecture itself was not significant (p=0.767 between HP configs).'
)

doc.add_heading('5.6 v7: 30-Seed Definitive Experiment', level=2)

add_para(
    'This is the definitive experiment: Deep [128, 128, 64] architecture, 30 seeds, dollar_delta reward, '
    'z-score normalization with 25x NLP upscale. Of 750 planned runs, 745 completed successfully (5 SVD '
    'convergence errors). The overall active rate was 70.9%: 71% active, 27% do-nothing, 2% oscillating. '
    'By ticker, active rates ranged from 47% (CRM) to 95% (NFLX).'
)

add_para(
    'The ensemble returns showed a 4.9 percentage point spread, but ANOVA yielded F=0.034, p=0.998, '
    'eta-squared=0.006. Even with maximum statistical power (750 runs, 30 seeds), no LLM effect was detected. '
    'The detailed results are presented in Section 6.'
)

doc.add_heading('5.7 v8: Ablation Study', level=2)

add_para(
    'Version 8 changed the question from "which LLM" to "which type of information matters." We compared three '
    'feature sets: Full (8 tech + 7 NLP), Tech-only (8 tech + 0 NLP), and NLP-only (0 tech + 7 NLP). This '
    'addressed a different question -- not whether LLM choice matters, but whether NLP information itself matters '
    'when compared to technical information. The results revealed a paradox that is resolved in Section 7.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. RESULTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('6. Results', level=1)

doc.add_heading('6.1 Stage 1: NLP Signal Quality', level=2)

add_para(
    'We first benchmarked the direct predictive quality of NLP signals generated by each LLM, independent of '
    'the trading system.'
)

add_table(
    ['LLM Configuration', 'Directional Accuracy', 'Information Coefficient', 'All-Zero-Day Rate', 'Per-Cell Zero-Fill'],
    [
        ['Gemini 3.1 Pro', '38.6%', '+0.024', '6.9%', '32.2%'],
        ['Qwen Base', '39.0%', '+0.025', '39.9%', '57.9%'],
        ['QLoRA', '41.1%', '+0.025', '18.2%', '82.6%'],
        ['QA-LoRA', '40.2%', '+0.009', '12.0%', '90.9%'],
        ['Random Baseline', '33.3%', '0.000', '\u2014', '\u2014'],
        ['Useful Threshold', '>50%', '>0.05', '\u2014', '\u2014'],
    ],
    col_widths=[3.0, 3.0, 3.0, 3.0, 3.0],
    caption='Stage 1 NLP signal quality metrics. Directional accuracy uses 3-class classification (up/down/neutral with \u00b10.5% threshold); random baseline is 33.3%. IC is Spearman rank correlation between raw NLP feature value and next-day return.'
)

add_para(
    'All four LLM configurations produce signals with directional accuracy barely above the 33.3% random baseline '
    '(range: 38.6\u201341.1%), and information coefficients well below the 0.05 threshold considered useful in '
    'quantitative finance (range: +0.009 to +0.025). Critically, no LLM significantly outperforms any other on '
    'these metrics (all pairwise p>0.99). Despite dramatic differences in zero-fill rates and signal distributions, '
    'the actual predictive content is comparably weak across all configurations.'
)

add_para(
    'The information coefficient (IC) is computed as the Spearman rank correlation between the raw NLP feature '
    'value and the next-day return, which is the standard definition in quantitative finance. Unlike Pearson '
    'correlation, Spearman measures monotonic relationship only and is robust to outliers. An IC above 0.05 is '
    'generally considered "useful" and above 0.10 is "strong" in the industry.'
)

doc.add_heading('6.2 Stage 2: Trading Performance (v7)', level=2)

doc.add_heading('6.2.1 Ensemble Returns', level=3)

add_table(
    ['Configuration', 'AAPL', 'AMZN', 'CRM', 'MSFT', 'NFLX', 'Mean', 'Sharpe'],
    [
        ['Gemini', '+16.2%', '+32.8%', '+28.7%', '-2.0%', '+61.6%', '+27.5%', '1.47'],
        ['Qwen Base', '+14.9%', '+27.2%', '+28.7%', '-2.0%', '+61.5%', '+26.1%', '1.38'],
        ['QLoRA', '+16.2%', '+32.8%', '+28.7%', '-1.9%', '+58.4%', '+26.8%', '1.44'],
        ['QA-LoRA', '+9.6%', '+27.4%', '+17.6%', '-1.7%', '+60.0%', '+22.6%', '1.19'],
        ['No-NLP', '+9.6%', '+27.4%', '+28.7%', '-2.0%', '+60.9%', '+24.9%', '1.30'],
    ],
    col_widths=[2.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.5],
    caption='Ensemble returns by configuration and ticker (v7, 30 seeds). ANOVA: F=0.034, p=0.998, \u03b7\u00b2=0.006.'
)

add_para(
    'The ensemble returns show a 4.9 percentage point spread: Gemini leads at +27.5%, QA-LoRA trails at +22.6%. '
    'However, ANOVA yields F=0.034, p=0.998, with eta-squared of just 0.006, meaning LLM choice explains less '
    'than 1% of return variance. The visual spread is misleading because each ensemble return is based on only '
    '5 tickers (far too few for inference) and within-group variance is enormous (MSFT near -2%, NFLX above +58%). '
    'The between-LLM variation is completely dominated by between-ticker variation.'
)

doc.add_heading('6.2.2 Individual Agent Level Analysis', level=3)

add_para(
    'The properly powered analysis uses paired same-seed comparisons at the individual agent level. By matching '
    'agents on (ticker, seed) pairs -- so that Gemini-MSFT-seed42 is compared only with NoNLP-MSFT-seed42 -- '
    'ticker and seed variation are differenced out. With 30 seeds per configuration, this yields 95\u2013100 '
    'paired observations per LLM comparison.'
)

add_table(
    ['Comparison', 'n (paired)', 'Difference (pp)', 'Raw p-value', 'BH-corrected p', 'Significant at 0.05?'],
    [
        ['Gemini vs No-NLP', '98', '+0.94', '0.033', '0.066', 'No'],
        ['Qwen Base vs No-NLP', '95', '+1.05', '0.058', '0.077', 'No'],
        ['QLoRA vs No-NLP', '100', '+0.73', '0.011', '0.044', 'Marginal'],
        ['QA-LoRA vs No-NLP', '98', '+0.04', '0.890', '0.890', 'No'],
        ['Pooled Has-NLP vs No-NLP', '\u2014', '+0.81', '0.736', '\u2014', 'No'],
    ],
    col_widths=[3.5, 1.8, 2.2, 2.0, 2.2, 2.8],
    caption='Paired same-seed t-tests (LLM vs No-NLP). BH = Benjamini-Hochberg correction for 4 comparisons. Pooled test uses unpaired comparison.'
)

add_para(
    'At the individual agent level, the true LLM vs No-NLP effects shrink dramatically. The largest effect is '
    'Qwen Base at +1.05pp; the smallest is QA-LoRA at +0.04pp. After Benjamini-Hochberg correction, only QLoRA '
    'achieves marginal significance with a tiny +0.73pp effect. The pooled Has-NLP vs No-NLP comparison yields '
    '+0.81pp, p=0.736, Cohen\'s d=0.037 -- a negligible effect size where return distributions overlap by 98%.'
)

doc.add_heading('6.3 Feature Importance Analysis', level=2)

add_para(
    'To understand why NLP signals have so little impact, we conducted permutation importance analysis on trained '
    'PPO policies. The importance of each feature is measured by the mean absolute change in the agent\'s action '
    'when that feature is randomly permuted.'
)

add_table(
    ['Configuration', 'Cash/Price/Holdings', 'Technical Indicators', 'NLP Features'],
    [
        ['Gemini-AAPL', '83.9%', '15.3%', '0.8%'],
        ['Gemini-MSFT', '88.8%', '11.0%', '0.2%'],
        ['No-NLP-AAPL', '85.4%', '14.6%', '0.0%'],
        ['Expected (uniform)', '~17%', '~44%', '~39%'],
    ],
    col_widths=[3.5, 3.5, 3.5, 3.0],
    caption='Permutation importance by feature group. Expected values assume equal importance per feature (3 portfolio / 8 tech / 7 NLP of 18 total).'
)

add_para(
    'Portfolio state variables (cash, price, holdings) dominate decision variance at 84\u201389%, while NLP '
    'features contribute only 0.2\u20130.8% -- far below the ~39% expected if all features were equally important. '
    'This is mechanically sensible: PPO\'s action is a buy/sell amount, and the hard constraints of available '
    'cash, current holdings, and price literally define the feasible action set. Per-feature, each portfolio '
    'variable carries approximately 22% importance versus ~4% per technical feature and ~0.2% per NLP feature. '
    'Cash alone is roughly 101 times more important than the average NLP feature.'
)

doc.add_heading('6.4 Ablation Results (v8)', level=2)

add_para(
    'The ablation study (v8) compared three feature sets -- Full (tech+NLP), Tech-only (NLP zeroed), and '
    'NLP-only (tech zeroed) -- to determine which type of information drives trading performance. This addresses '
    'a different question than the LLM comparison: not which LLM, but which type of information matters.'
)

doc.add_heading('6.4.1 Active Rates', level=3)

add_table(
    ['Configuration', 'Full Active/Total', 'Full Rate', 'NLP-Only Active/Total', 'NLP-Only Rate'],
    [
        ['Gemini', '109/150', '72.7%', '145/150', '96.7%'],
        ['Qwen Base', '103/148', '69.6%', '145/150', '96.7%'],
        ['QLoRA', '107/149', '71.8%', '148/150', '98.7%'],
        ['QA-LoRA', '106/150', '70.7%', '146/150', '97.3%'],
        ['Tech-Only', '103/148', '69.6%', '\u2014', '\u2014'],
    ],
    col_widths=[2.5, 3.0, 2.0, 3.5, 2.0],
    caption='Active agent rates by feature set. NLP-only agents are consistently more active than Full or Tech-only agents.'
)

add_para(
    'The most striking finding is the active rate gap: NLP-only agents are 96.7\u201398.7% active, compared to '
    'only 69.6% for Tech-only agents -- a 27 percentage point difference. PPO with correlated technical features '
    'falls into do-nothing traps far more often. Additionally, NLP-only agents that do trade make 3\u20136 trades '
    '(decisive, buy-and-hold-like), while Tech-only agents make 7\u201311 trades (erratic, suggesting difficulty '
    'converging to a stable policy).'
)

doc.add_heading('6.4.2 Paired Comparisons', level=3)

add_table(
    ['Comparison', 'Difference (pp)', 't-statistic', 'p-value', 'Cohen\'s d', 'Significance'],
    [
        ['NLP-only vs Tech-only', '+1.8', '5.51', '<0.0001', '0.36', '***'],
        ['NLP-only vs Full', '+1.3', '4.08', '0.0001', '0.26', '***'],
        ['Full vs Tech-only', '+0.5', '2.63', '0.009', '0.12', '**'],
    ],
    col_widths=[3.5, 2.5, 2.0, 2.5, 2.0, 2.0],
    caption='Ablation paired t-tests (pooled across LLM configs, excluding NFLX). **p<0.01, ***p<0.001.'
)

add_para(
    'Every comparison is statistically significant: NLP-only outperforms both Full and Tech-only, and Full '
    'outperforms Tech-only. This holds across 3 of 4 LLM configurations individually (Qwen Base is borderline '
    'at p=0.051). The per-configuration breakdown is as follows:'
)

add_table(
    ['LLM Config', 'NLP vs Tech (pp)', 'p-value', 'NLP vs Full (pp)', 'p-value', 'Full vs Tech (pp)', 'p-value'],
    [
        ['Gemini', '+2.05', '0.002', '+1.56', '0.011', '+0.49', '0.223'],
        ['Qwen Base', '+1.40', '0.051', '+0.50', '0.435', '+0.89', '0.124'],
        ['QLoRA', '+2.04', '0.006', '+1.33', '0.073', '+0.71', '0.038'],
        ['QA-LoRA', '+1.77', '0.003', '+1.72', '0.002', '+0.05', '0.850'],
    ],
    col_widths=[2.2, 2.5, 1.8, 2.5, 1.8, 2.5, 1.8],
    caption='Per-LLM configuration ablation results (excluding NFLX). Bold p-values <0.05.'
)

add_para(
    'This creates an apparent paradox: NLP-only beats Tech-only by +1.8pp (p<0.0001), yet in the v7 experiment, '
    'adding NLP to tech makes zero difference (Has-NLP vs No-NLP: +0.81pp, p=0.736). NLP appears both useful '
    '(when alone) and useless (when added to tech). The resolution is presented in Section 7.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. DISCUSSION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('7. Discussion', level=1)

doc.add_heading('7.1 Why NLP Signals Do Not Improve Trading', level=2)

add_para(
    'The persistence of the null result across 8 experimental versions -- including fixes for feature scale, '
    'reward function, architecture, and sample size -- suggests a fundamental explanation. We identify six '
    'contributing factors:'
)

add_bullet('In the raw observation space (v1/v2), NLP features were orders of magnitude smaller than portfolio '
           'state variables, making them invisible to gradient-based learning. This was fixed in v3+ but the null persists.', bold_prefix='Feature scale mismatch: ')
add_bullet('The 8 technical indicators plus closing price provide sufficient signal for the PPO agent to learn '
           'buy-and-hold strategies in trending markets. NLP adds no marginal information because it is redundant '
           'with technical indicators.', bold_prefix='Information redundancy: ')
add_bullet('The median agent makes only 3.2 trades per 144-day test period (\u224844-day holding period), behaving '
           'essentially as buy-and-hold. Once a position is established, daily sentiment updates do not alter it.', bold_prefix='Buy-and-hold convergence: ')
add_bullet('Approximately 27% of agents converge to do-nothing policies regardless of LLM configuration, and '
           'this bimodal distribution (active vs. inactive) overwhelms any subtle NLP effect.', bold_prefix='Do-nothing seed dominance: ')
add_bullet('NLP signals are generated daily but market reactions may be intraday or delayed by uncertain amounts, '
           'creating a temporal misalignment between signal and price impact.', bold_prefix='Temporal mismatch: ')
add_bullet('A zero value in NLP features conflates "no news" with "neutral sentiment," an unresolvable ambiguity. '
           'Even Gemini (6.9% all-zero-day rate) shows this is a fundamental signal quality limitation, not a '
           'coverage issue.', bold_prefix='Zero-fill ambiguity: ')

doc.add_heading('7.2 The Ablation Paradox and Resolution', level=2)

add_para(
    'The ablation study (v8) reveals an apparent paradox: NLP-only agents significantly outperform Tech-only '
    'agents (+1.8pp, p<0.0001), yet adding NLP to tech produces no improvement (+0.81pp, p=0.736). How can NLP '
    'be both useful and useless?'
)

add_para(
    'The resolution is that NLP and technical indicators encode similar information -- they are redundant -- but '
    'they deliver this information in different forms. The ablation itself proves the redundancy: if NLP and tech '
    'carried different information, removing one would hurt significantly more than we observe. Instead, '
    'NLP-only \u2248 Tech-only \u2248 Full in terms of what the agent can learn.'
)

add_para(
    'The critical difference is the form of delivery:'
)

add_bullet('7 features, less correlated with each other. Active rate: 96.7\u201398.7%. Trade count: 3\u20136 (decisive).', bold_prefix='NLP-only: ')
add_bullet('8 features, highly correlated (CCI, Bollinger Bands, RSI, MACD all measure momentum and trend variations). '
           'Active rate: 69.6%. Trade count: 7\u201311 (erratic).', bold_prefix='Tech-only: ')
add_bullet('15 features, redundant information. Active rate: 69.6\u201372.7%. Trade count: 3\u20136.', bold_prefix='Full (Tech+NLP): ')

add_para(
    'The 27 percentage point active rate gap (97.3% vs. 69.6%) is the key mechanism. PPO with correlated '
    'technical features frequently falls into do-nothing traps, where the gradient landscape has flat regions '
    'that the agent cannot escape. The simpler, less correlated NLP state space produces a smoother optimization '
    'landscape that helps PPO converge to active policies.'
)

add_para(
    'Thus, the two results are not contradictory:'
)

add_bullet('Adding NLP to tech = same information twice = no gain. This explains the null result.', bold_prefix='Redundancy: ')
add_bullet('Replacing tech with NLP = same information, simpler form = easier learning. This explains the ablation result.', bold_prefix='Simpler form: ')

add_para(
    'In summary: NLP \u2248 Tech in information content (redundancy, not irrelevance). The null result holds because '
    'NLP adds no new information to what tech already provides. But NLP alone suffices for effective learning '
    'because its simpler state space facilitates PPO optimization.'
)

doc.add_heading('7.3 Ensemble Granularity', level=2)

add_para(
    'The ensemble returns in Table 4 show a 4.9pp spread that visually suggests a meaningful LLM effect, yet '
    'ANOVA declares it insignificant. This discrepancy arises from the ensemble construction methodology. Each '
    'ensemble return is based on only 5 tickers (n=5), which is far too few for statistical inference. The '
    'properly powered test at the individual agent level (n=95\u2013100) reveals the true effect sizes are under 1pp.'
)

add_para(
    'Furthermore, the ensemble\'s action-averaging procedure creates a nonlinear mapping from small action '
    'differences to large return differences. As evidence: the AAPL ensemble gap between Gemini and No-NLP is '
    '+6.6pp, but the individual agent gap is only +1.0pp -- a 6x amplification. The ensemble magnifies small '
    'differences through aggregation nonlinearity, not because the underlying effect is large.'
)

doc.add_heading('7.4 Alternative Algorithms', level=2)

add_para(
    'We tested SAC (Soft Actor-Critic) extensively as an alternative to PPO. SAC converged to pure buy-and-hold '
    'behavior in 49 out of 50 runs, with returns that were byte-for-byte identical across LLM configurations. '
    'SAC\'s entropy regularization appears to suppress the exploration needed for active trading in this environment. '
    'TD3 produced similar patterns. PPO was the only algorithm that produced active trading behavior, and the null '
    'result holds for it as well. This suggests the null is fundamental to the information structure rather than '
    'algorithm-specific.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('8. Conclusion', level=1)

add_para(
    'This study set out to answer a straightforward question: does better NLP signal quality, as produced by '
    'more capable LLMs, improve the trading performance of a DRL agent? After 8 experimental iterations and '
    'over 2,800 training runs, the answer is unequivocally no.'
)

add_para(
    'Three key findings emerge:'
)

add_rich_para([
    ('1. LLM choice has no significant effect on trading performance. ', True, False, 11, NAVY),
    ('One-way ANOVA across the definitive 30-seed experiment (750 runs) yields F=0.034, p=0.998, \u03b7\u00b2=0.006. '
     'Even the broadest comparison -- having NLP features versus not -- shows only +0.81pp difference (p=0.736, '
     'Cohen\'s d=0.037), with 98% overlap in return distributions. This null result persists across reward functions, '
     'architectures, and normalization strategies.', False, False, 11, DARK),
], indent=0.5, space_after=8)

add_rich_para([
    ('2. NLP-only agents outperform Tech-only agents by +1.8pp (p<0.0001, d=0.36). ', True, False, 11, NAVY),
    ('This is not because NLP carries stronger predictive signals -- directional accuracy is only 38\u201341% '
     'and IC is 0.01\u20130.03, both near random. Rather, the 7 NLP features are less correlated than the 8 '
     'technical features, creating a simpler state space that helps PPO avoid do-nothing traps (97.3% active '
     'vs. 69.6%).', False, False, 11, DARK),
], indent=0.5, space_after=8)

add_rich_para([
    ('3. NLP \u2248 Tech in information content: redundancy, not irrelevance. ', True, False, 11, NAVY),
    ('The ablation study proves that NLP and tech encode similar information: removing either one leaves the '
     'agent with sufficient signal for effective trading. Adding NLP to tech is redundant (no gain), but NLP '
     'alone is sufficient because its simpler form facilitates learning.', False, False, 11, DARK),
], indent=0.5, space_after=8)

add_para('Practical Implications', bold=True, size=12, color=NAVY, space_after=6)

add_bullet('Investing in better LLMs for trading signal generation yields zero marginal return, because '
           'technical indicators already capture the same information.', bold_prefix='LLM selection: ')
add_bullet('Fewer, less correlated features may outperform more features. The 27pp active rate gap between '
           'NLP-only (97%) and Tech-only (70%) demonstrates that feature correlation matters more than feature '
           'count for PPO learning.', bold_prefix='State space design: ')
add_bullet('LLM quality differences should be evaluated at the NLP signal level (accuracy, F1, agreement), '
           'not at the trading return level -- trading is insensitive to NLP quality.', bold_prefix='Evaluation methodology: ')
add_bullet('The finding that RL is insensitive to LLM quality is itself a valid and important result. It '
           'challenges the common assumption that better NLP naturally leads to better trading.', bold_prefix='Research contribution: ')

add_para('Limitations and Future Work', bold=True, size=12, color=NAVY, space_after=6)

add_bullet('Our 5 test tickers are all large-cap US technology stocks that appreciated during the test period, '
           'introducing survivorship and selection bias. Testing on a broader universe including bear markets '
           'would strengthen the conclusions.', bold_prefix='Ticker selection: ')
add_bullet('The 7 NLP signals from PrimoGPT may not capture all dimensions of news impact. Alternative signal '
           'designs (e.g., fine-grained sentiment, event extraction, topic modeling) could produce different results.', bold_prefix='Signal design: ')
add_bullet('Our PPO agent learns buy-and-hold-like strategies with a median of 3.2 trades per 144 days. In a '
           'higher-frequency or more active trading regime, NLP signals might play a larger role.', bold_prefix='Trading frequency: ')
add_bullet('The cross-LLM mean signal correlation is approximately 0.6, and 66\u201396% of non-zero cells differ. '
           'Despite this apparent diversity, the signals converge to similar downstream effects, suggesting the '
           'differences are in noise rather than signal.', bold_prefix='Cross-LLM agreement: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. REFERENCES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('9. References', level=1)

refs = [
    'Botunac, L. (2025). Automated Trading Framework Using LLM-Driven Features and Deep Reinforcement Learning. Big Data and Cognitive Computing, 9(3), 67.',
    'Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). Proximal Policy Optimization Algorithms. arXiv preprint arXiv:1707.06347.',
    'Haarnoja, T., Zhou, A., Abbeel, P., & Levine, S. (2018). Soft Actor-Critic: Off-Policy Maximum Entropy Deep Reinforcement Learning with a Stochastic Actor. Proceedings of ICML.',
    'Hu, E. J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., ... & Chen, W. (2021). LoRA: Low-Rank Adaptation of Large Language Models. arXiv preprint arXiv:2106.09685.',
    'Dettmers, T., Pagnoni, A., Holtzman, A., & Zettlemoyer, L. (2024). QLoRA: Efficient Finetuning of Quantized Large Language Models. Proceedings of NeurIPS.',
    'Renduchintala, A., Li, T., & Gao, J. (2024). QA-LoRA: Quantization-Aware Low-Rank Adaptation of Large Language Models. arXiv preprint arXiv:2309.14717.',
    'Yang, H., Liu, X. Y., & Wang, C. D. (2020). Deep Reinforcement Learning for Automated Stock Trading: An Ensemble Strategy. Proceedings of ACM ICAIF.',
    'Benjamini, Y., & Hochberg, Y. (1995). Controlling the False Discovery Rate: A Practical and Powerful Approach to Multiple Testing. Journal of the Royal Statistical Society, Series B, 57(1), 289\u2013300.',
    'Sorensen, E. H., Hua, R., & Qian, E. (2022). Contextual Synergies Between Large Language Models and Quantitative Investing. Journal of Financial Data Science, 4(4), 102\u2013118.',
    'Liu, X., Wang, Y., & Chen, H. (2023). FinGPT: Open-Source Financial Large Language Models. Proceedings of ACM ICAIF.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}] ')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run = p.add_run(ref)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Appendix A: Hyperparameters and Seed List', level=1)

doc.add_heading('A.1 PPO Hyperparameters (v7 Definitive Experiment)', level=2)

add_table(
    ['Parameter', 'Value', 'Notes'],
    [
        ['Architecture', '[128, 128, 64]', 'Deep configuration'],
        ['Learning Rate', '1e-4', 'Adam optimizer'],
        ['Batch Size', '256', 'Mini-batch size'],
        ['N-Steps', '2048', 'Rollout buffer length'],
        ['Entropy Coefficient', '0.01', 'Encourages exploration'],
        ['Clip Range', '\u00b10.2', 'PPO clip parameter'],
        ['Total Timesteps', '200,000', 'Per training run'],
        ['Gamma', '0.99', 'Discount factor'],
        ['GAE Lambda', '0.95', 'Generalized advantage estimation'],
        ['Reward Function', 'dollar_delta', 'Change in portfolio value'],
        ['Normalization', 'z-score + 25x NLP', 'On cash/price/holdings and NLP features'],
        ['Transaction Cost', '0.1%', 'Per trade'],
        ['Margin Call', '25%', 'Stop if equity < 25% of initial'],
    ],
    col_widths=[3.5, 3.0, 6.0],
    caption='Complete PPO hyperparameters for the v7 definitive experiment'
)

doc.add_heading('A.2 Seed List', level=2)

add_para(
    'The 30 seeds used in v7 and v8 experiments: 42, 123, 456, 789, 2024, 314, 271, 1618, 999, 7, '
    '100, 200, 300, 400, 500, 600, 700, 800, 900, 1000, 1111, 2222, 3333, 4444, 5555, 6666, 7777, '
    '8888, 9999, 12345.'
)

doc.add_heading('A.3 LLM Fine-Tuning Details', level=2)

add_table(
    ['Configuration', 'Base Model', 'Method', 'Rank (r)', 'Quantization', 'Training Data'],
    [
        ['QLoRA', 'Qwen3.5-27B', 'QLoRA', '16', '4-bit NF4', 'Financial news 2022\u20132024'],
        ['QA-LoRA', 'Qwen3.5-27B', 'QA-LoRA', '16', '4-bit', 'Financial news 2022\u20132024'],
    ],
    col_widths=[2.5, 2.5, 2.0, 2.0, 2.5, 4.0],
    caption='LLM fine-tuning configurations'
)

doc.add_heading('A.4 Data Pipeline', level=2)

add_para(
    'Market data is sourced via yfinance (auto_adjust=True) for the 5 test tickers (AAPL, AMZN, CRM, MSFT, NFLX) '
    'over the training period (2022-04 to 2024-07) and test period (2024-08 to 2025-02). Technical indicators '
    'are computed using the FinRL feature engineering pipeline. NLP signals are pre-generated by each LLM and '
    'merged on (ticker, date) keys, then shifted by one day to prevent look-ahead bias. The training set uses '
    'different reference tickers (META, GOOGL, TSLA, IBM, AMD) as in the original study.'
)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Research_Report.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
